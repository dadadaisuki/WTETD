from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "six-feature-guide.docx"


def set_run_font(run, name="Arial", size=11, bold=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_style(paragraph, before=0, after=8, line_spacing=1.15, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    if align is not None:
        paragraph.alignment = align


def add_paragraph(doc, text, *, size=11, bold=False, color="000000", before=0, after=8, align=None):
    paragraph = doc.add_paragraph()
    set_paragraph_style(paragraph, before=before, after=after, line_spacing=1.15, align=align)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    size = 16 if level == 1 else 13
    before = 18 if level == 1 else 12
    after = 6 if level == 1 else 4
    return add_paragraph(doc, text, size=size, bold=False, before=before, after=after)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_style(paragraph, before=0, after=4, line_spacing=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    set_paragraph_style(paragraph, before=0, after=4, line_spacing=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def read_lines(relative_path, start, end):
    path = ROOT / relative_path
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    rows = []
    for idx in range(start, end + 1):
        if 1 <= idx <= len(lines):
            rows.append((idx, lines[idx - 1]))
    return rows


def add_code_excerpt(doc, title, relative_path, start, end):
    add_paragraph(doc, f"代码摘录：{title}", bold=True, before=4, after=6)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    header = table.rows[0].cells
    header[0].width = Inches(0.7)
    header[1].width = Inches(5.8)
    header[0].text = "行号"
    header[1].text = f"{relative_path}:{start}-{end}"

    for cell in header:
        for paragraph in cell.paragraphs:
            set_paragraph_style(paragraph, before=0, after=0, line_spacing=1.0)
            for run in paragraph.runs:
                set_run_font(run, size=9, bold=True)

    for line_no, code in read_lines(relative_path, start, end):
        row = table.add_row().cells
        row[0].width = Inches(0.7)
        row[1].width = Inches(5.8)
        row[0].text = str(line_no)
        row[1].text = code or " "
        for cell in row:
            for paragraph in cell.paragraphs:
                set_paragraph_style(paragraph, before=0, after=0, line_spacing=1.0)
                for run in paragraph.runs:
                    set_run_font(run, name="Consolas", size=8)

    add_paragraph(doc, f"源码定位：{relative_path}:{start}-{end}", size=10, color="555555", before=4, after=8)


def add_line_explanations(doc, items):
    add_paragraph(doc, "逐行讲解：", bold=True, before=4, after=6)
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        set_paragraph_style(paragraph, before=0, after=4, line_spacing=1.15)
        prefix = paragraph.add_run(item["line"])
        set_run_font(prefix, size=11, bold=True)
        body = paragraph.add_run(item["explain"])
        set_run_font(body, size=11)


def add_feature(doc, feature):
    add_heading(doc, feature["title"], level=1)
    add_paragraph(doc, f"对应考核要求：{feature['requirement']}", bold=True, before=0, after=4)
    add_paragraph(doc, f"功能定位：{feature['goal']}", bold=True, before=0, after=6)
    add_paragraph(doc, feature["summary"])
    add_paragraph(doc, "主要文件与函数：", bold=True, before=4, after=4)
    for bullet in feature["functions"]:
        add_bullet(doc, bullet)
    for excerpt in feature["excerpts"]:
        add_code_excerpt(doc, excerpt["title"], excerpt["path"], excerpt["start"], excerpt["end"])
    add_line_explanations(doc, feature["explanations"])
    add_paragraph(doc, f"答辩时可以这样讲：{feature['defense']}", bold=True, before=6, after=8)


FEATURES = [
    {
        "title": "功能 1：顶层视图切换与动态组件渲染",
        "requirement": "要求（1）动态数据渲染与视图切换",
        "goal": "不使用路由插件，只靠响应式状态完成单页应用的页面切换。",
        "summary": "项目最顶层在 App.vue 中维护 currentView 这个响应式状态，再把页面 key 和具体组件放进 viewMap 里。这样一来，只要 currentView 发生变化，<component :is=\"...\"> 就会自动切换当前显示的页面，完全符合 Vue 的“数据驱动视图”思想。",
        "functions": [
            "src/App.vue 中的 currentView、viewMap、currentViewComponent、handleMenuChange、ensureViewData。",
            "src/components/AppModuleNav.vue 中的 defineEmits(['menu-change']) 和 changeMenu。",
        ],
        "excerpts": [
            {"title": "顶层状态与视图切换核心逻辑", "path": "src/App.vue", "start": 19, "end": 71},
            {"title": "视图切换后的生命周期补充加载", "path": "src/App.vue", "start": 84, "end": 95},
            {"title": "模板中的动态组件与事件绑定", "path": "src/App.vue", "start": 125, "end": 137},
            {"title": "导航组件向上抛出切换事件", "path": "src/components/AppModuleNav.vue", "start": 1, "end": 35},
        ],
        "explanations": [
            {"line": "第 19 行：", "explain": "const currentView = ref('home') 定义了当前页面的总状态。它是一个 ref，所以模板里会自动解包，页面切换时不需要你手动写 .value。"},
            {"line": "第 46 到 51 行：", "explain": "viewMap 把字符串 key 和实际组件做映射。这个设计的优点是结构清晰、后期新增页面扩展很方便。"},
            {"line": "第 53 到 55 行：", "explain": "currentViewComponent 是 computed。它会根据当前状态动态返回应该渲染哪个组件，如果 key 不存在则兜底回到首页。"},
            {"line": "第 61 到 67 行：", "explain": "handleMenuChange 不直接操作 DOM，而是只改 currentView。Vue 发现状态变了，就会自动更新界面。"},
            {"line": "第 69 到 71 行：", "explain": "ensureViewData 在切换视图的同时补充加载该页面需要的数据，避免一次性把所有接口都打满。"},
            {"line": "第 90 到 92 行：", "explain": "watch(currentView, ...) 说明只要页面 key 改变，就会再次触发数据加载，这也是“视图切换”和“懒加载数据”之间的连接点。"},
            {"line": "第 132 到 137 行：", "explain": "<component :is=\"currentViewComponent\"> 是整套无路由切页的关键。它渲染的是一个“变量代表的组件”，而不是写死的标签。"},
            {"line": "导航组件第 9 到 20 行：", "explain": "导航组件通过 defineEmits 声明 menu-change 事件，再把 nextView 上抛给父组件。子组件不保留总状态，父组件统一管理，符合单向数据流。"},
        ],
        "defense": "我用 currentView 管理当前页面，用 viewMap 维护页面映射，再通过动态组件 <component :is=\"...\"> 完成页面切换。整个过程没有用原生 DOM 控制显隐，而是纯粹依赖 Vue 的响应式状态。",
    },
    {
        "title": "功能 2：看板与轮盘候选池的动态筛选渲染",
        "requirement": "要求（1）动态数据渲染与视图切换",
        "goal": "让店铺列表和轮盘候选池都能根据用户输入实时重算并更新界面。",
        "summary": "这个功能把“动态渲染”做得更深入：推荐看板里，商家列表会随着搜索词、Tag 和展开状态变化；轮盘页里，候选菜品池会随着 Tag、食材和搜索词变化，并且概率和扇区面积也会随之重算。也就是说，用户看到的不是固定页面，而是实时推导出来的界面结果。",
        "functions": [
            "src/views/DashboardView.vue 中的 filteredMerchants、dashboardTags、getVisibleMerchantDishes、toggleMerchant。",
            "src/views/WheelView.vue 中的 candidateDishes、calculateDishWeight、weightedDishes。",
            "src/components/TagWheel.vue 中的 totalWeight、sectors、wheelStyle。",
        ],
        "excerpts": [
            {"title": "推荐看板中的搜索、筛选与卡片展开", "path": "src/views/DashboardView.vue", "start": 118, "end": 188},
            {"title": "推荐看板模板中的 Tag 抽屉与商家列表", "path": "src/views/DashboardView.vue", "start": 200, "end": 327},
            {"title": "轮盘页中的候选池过滤与概率计算", "path": "src/views/WheelView.vue", "start": 140, "end": 180},
            {"title": "轮盘组件中的扇区角度与样式推导", "path": "src/components/TagWheel.vue", "start": 42, "end": 91},
        ],
        "explanations": [
            {"line": "Dashboard 第 118 到 145 行：", "explain": "filteredMerchants 是推荐看板的核心 computed。它会把 merchants、dishes、selectedTags 和 searchQuery 组合起来，得到当前应该显示的店铺列表。"},
            {"line": "第 147 到 149 行：", "explain": "dashboardTags 把场景标签和菜品标签合并去重，自动生成筛选抽屉的选项，所以筛选项本身也是数据推导结果。"},
            {"line": "第 169 到 176 行：", "explain": "watch(filteredMerchants, ...) 是一个细节优化。筛选条件变化后，如果之前展开的店铺已经不在结果里，就会自动收起，避免界面状态错乱。"},
            {"line": "模板第 210 到 233 行：", "explain": "Tag 抽屉用 v-show 控制开合，用 checkbox 配合 selectedTags 控制筛选状态。这正是“数据一变，列表就变”的声明式渲染。"},
            {"line": "模板第 236 到 323 行：", "explain": "商家卡片列表通过 v-for 渲染 filteredMerchants，再用 v-if 和 expandedMerchantId 控制详情区展开，形成手风琴式交互。"},
            {"line": "WheelView 第 140 到 157 行：", "explain": "candidateDishes 会按 Tag、食材和搜索词去过滤 dishes。轮盘候选池不是固定数组，而是会随着条件变化实时更新。"},
            {"line": "第 159 到 180 行：", "explain": "calculateDishWeight 和 weightedDishes 会把每个菜品换算成权重和概率，让轮盘不是平均八等分，而是根据数据动态改变扇区占比。"},
            {"line": "TagWheel 第 46 到 91 行：", "explain": "totalWeight、sectors 和 wheelStyle 把“候选数据”进一步翻译成“扇区角度”和“渐变样式”，最后交给模板自动渲染成轮盘。"},
        ],
        "defense": "这个功能说明我不只是做了页面切换，还做了更细粒度的动态渲染。推荐看板和轮盘页都不是静态展示，而是由 computed 和响应式状态实时推导出来的。",
    },
    {
        "title": "功能 3：交互式数据操作（新增、修改与热度更新）",
        "requirement": "要求（2）交互式数据操作（增删改查）",
        "goal": "把页面上的新增店铺、新增菜品、协同 Tag 更新和热度累加统一收口到数据层处理。",
        "summary": "项目里最关键的交互式数据操作不在页面模板里，而是在 useDiningStore 这个数据容器里。页面只负责收集输入和触发事件，真正的数据变更由 addMerchant、addDish、updateMerchantTags、incrementDishHeat 等函数统一完成，再决定是否同步到云端。",
        "functions": [
            "src/composables/useDiningStore.js 中的 addMerchant、addDish、updateMerchantTags、addMerchantTag、incrementDishHeat。",
            "src/views/WheelView.vue 中的 confirmResult。",
            "src/services/supabaseGateway.js 中预留的 deleteMerchant、deleteDish。",
        ],
        "excerpts": [
            {"title": "store 中的新增、修改与热度更新函数", "path": "src/composables/useDiningStore.js", "start": 292, "end": 460},
            {"title": "轮盘结果确认后触发热度更新", "path": "src/views/WheelView.vue", "start": 229, "end": 240},
            {"title": "前端 API 层中预留的删除接口", "path": "src/services/supabaseGateway.js", "start": 45, "end": 85},
        ],
        "explanations": [
            {"line": "第 292 到 316 行：", "explain": "addMerchant 会给新店铺生成 id、标准化名称与标签、补上默认热度和更新时间，然后把它插入 merchants 数组，再交给 queueSync 同步。"},
            {"line": "第 318 到 354 行：", "explain": "addDish 会先找到所属店铺，再结合店铺标签和菜品标签，判断是否自动补上“外卖”标签。这说明数据规则被放在了统一的数据层。"},
            {"line": "第 356 到 393 行：", "explain": "updateMerchantTags 和 addMerchantTag 负责店铺协同 Tag 的更新。这里不是直接改某个字段，而是通过 map 返回新的 merchants 数组，方便 Vue 正确追踪响应式变化。"},
            {"line": "第 405 到 460 行：", "explain": "incrementDishHeat 会在用户确认结果后给菜品 heat 加 1，同时给所属商家的 heat 也加 1。修改完成后，又会异步 patch 回云端接口。"},
            {"line": "WheelView 第 229 到 240 行：", "explain": "confirmResult 表示真正的数据修改时机是用户点击“就它了”，而不是轮盘一停就自动改热度。这样逻辑更合理，也更符合真实业务。"},
            {"line": "supabaseGateway 第 75 到 83 行：", "explain": "deleteMerchant 和 deleteDish 说明接口层已经预留了删除能力。也就是说，这个项目的数据层已经具备完整 CRUD 的基础。"},
        ],
        "defense": "这一项我会重点讲 store。页面本身不直接改全局数据，而是把新增、修改和热度更新都统一交给 useDiningStore 处理，这样逻辑集中、维护成本低，也更符合单向数据流。",
    },
    {
        "title": "功能 4：表单双向绑定与校验反馈",
        "requirement": "要求（3）表单双向绑定与验证反馈",
        "goal": "收集用户录入的店铺、菜品和自定义标签，并在前端就地拦截非法输入。",
        "summary": "管理页集中了项目里最完整的表单逻辑。这里既有 v-model.trim 绑定文本、数字和下拉框，也有 validateTag、validateMerchantForm、validateForm 这些防御性校验函数。用户新增协同 Tag 时，也会先校验再进入数据层处理。",
        "functions": [
            "src/views/ManageView.vue 中的 validateTag、validateMerchantForm、validateForm、submitMerchant、submitDish、addCollaborativeTag。",
            "src/components/FeedbackBar.vue 中的 message 条件渲染与动态 class。",
        ],
        "excerpts": [
            {"title": "管理页中的校验与提交函数", "path": "src/views/ManageView.vue", "start": 63, "end": 263},
            {"title": "管理页模板中的 v-model 绑定与反馈条", "path": "src/views/ManageView.vue", "start": 290, "end": 455},
            {"title": "反馈条组件", "path": "src/components/FeedbackBar.vue", "start": 1, "end": 18},
        ],
        "explanations": [
            {"line": "第 63 到 79 行：", "explain": "validateTag 是最基础的标签校验函数。它会依次检查：输入是否为空、是否超出长度限制、是否和已有标签重复。"},
            {"line": "第 81 到 103 行：", "explain": "addMerchantSceneTag 和 addMerchantCustomTag 都会先调用 validateTag。验证通过才 push 进数组，否则直接返回错误信息。"},
            {"line": "第 145 到 160 行：", "explain": "addCollaborativeTag 先校验再注入协同 Tag，这体现了“先拦截非法输入，再执行真正业务逻辑”的防御性编程思路。"},
            {"line": "第 162 到 183 行：", "explain": "validateMerchantForm 负责检查店铺名、位置说明和 Tag 数量，确保店铺表单提交前满足最基本要求。"},
            {"line": "第 213 到 247 行：", "explain": "validateForm 负责检查菜品名称、所属店铺、热量、价格、Tag 和食材是否合规，是菜品录入前的最后一道前端防线。"},
            {"line": "模板第 296、301、307、315、348、352、366、370 行：", "explain": "这些输入控件都通过 v-model 或 v-model.trim 和响应式状态绑定，所以输入变化会立刻同步到 merchantForm 或 form。"},
            {"line": "模板第 340、390、412、438、453 行：", "explain": "多个 FeedbackBar 说明页面把“校验逻辑”和“视觉提示”拆开了。方法只负责更新状态，组件只负责展示。"},
            {"line": "FeedbackBar 第 15 到 17 行：", "explain": "反馈条通过 v-if 判断是否显示，通过动态 class 切换 success、error、info 视觉样式，非常适合答辩时讲“声明式反馈”。"},
        ],
        "defense": "这个功能最适合答辩时讲防御性编程。我的表单不是提交后弹窗，而是通过 v-model 收集输入，再用校验函数拦截错误，最后由 FeedbackBar 自动把结果渲染出来。",
    },
    {
        "title": "功能 5：浏览器原生能力集成",
        "requirement": "要求（4）浏览器原生能力（BOM / Web API）集成",
        "goal": "利用 localStorage、定时器、震动和滚动 API 增强交互体验与状态记忆。",
        "summary": "这一项项目里实现得很完整：useDiningStore 会把核心数据写进 localStorage，WheelView 会把用户的筛选组合写进 localStorage，并使用 setInterval、setTimeout 和 navigator.vibrate 做冷却保护与震动反馈，App.vue 还用了 window.scrollTo 和滚动监听实现回到顶部按钮。",
        "functions": [
            "src/composables/useDiningStore.js 中的 readLocalData、writeLocalData。",
            "src/views/WheelView.vue 中的 persistWheelPreferences、rememberCurrentCombination、vibrateResult、startCooldown。",
            "src/App.vue 中的 handleWindowScroll、scrollToTop。",
        ],
        "excerpts": [
            {"title": "store 中的 localStorage 读写", "path": "src/composables/useDiningStore.js", "start": 68, "end": 101},
            {"title": "轮盘页中的筛选历史持久化", "path": "src/views/WheelView.vue", "start": 57, "end": 124},
            {"title": "轮盘页中的震动与冷却计时器", "path": "src/views/WheelView.vue", "start": 199, "end": 240},
            {"title": "顶层页面中的滚动监听与回到顶部", "path": "src/App.vue", "start": 73, "end": 95},
        ],
        "explanations": [
            {"line": "store 第 68 到 84 行：", "explain": "readLocalData 负责从 localStorage 安全读取缓存。它先判断 window 是否存在，再尝试 JSON.parse，避免非浏览器环境或坏缓存导致报错。"},
            {"line": "第 86 到 101 行：", "explain": "writeLocalData 会把 merchants、dishes 和 homeSnapshot 序列化进 localStorage，而 watch([...], writeLocalData, { deep: true }) 保证它们一变就自动保存。"},
            {"line": "WheelView 第 57 到 95 行：", "explain": "persistWheelPreferences 和 rememberCurrentCombination 会记录用户的 Tag、食材组合，并把最近历史写进 localStorage，属于非常典型的个性化偏好保存。"},
            {"line": "第 199 到 203 行：", "explain": "vibrateResult 使用 navigator.vibrate 调用浏览器原生震动能力。这里先做能力检测，再调用接口，是比较规范的 Web API 用法。"},
            {"line": "第 205 到 221 行：", "explain": "startCooldown 里同时用了 setInterval 和 setTimeout：前者负责每秒更新倒计时，后者负责 10 秒后自动解除冷却。"},
            {"line": "App.vue 第 73 到 81 行：", "explain": "handleWindowScroll 会根据 window.scrollY 控制 showBackToTop；scrollToTop 则调用 window.scrollTo({ top: 0, behavior: 'smooth' }) 平滑回顶。"},
            {"line": "第 84 到 95 行：", "explain": "onMounted 时注册滚动监听，onBeforeUnmount 时解除监听，说明这个功能不是随手写的，而是配合组件生命周期做了规范管理。"},
        ],
        "defense": "这一项我会重点讲 localStorage 和浏览器原生 API。项目不仅能保存用户偏好，还用定时器、滚动和震动能力增强了交互体验，完全符合老师要求的 BOM / Web API 集成。",
    },
    {
        "title": "功能 6：异步数据获取与展示",
        "requirement": "要求（5）异步数据获取与展示",
        "goal": "通过 fetch 从接口异步获取数据，再把数据写入响应式状态驱动页面展示。",
        "summary": "项目不是写死数据，而是通过自己的 /api 接口异步获取首页摘要、商家列表和菜品列表。最关键的地方在于：接口层先被统一封装为 supabaseGateway，再由 useDiningStore 的 loadCatalog 按页面范围懒加载数据，最后交给各个视图中的 computed 和模板渲染。",
        "functions": [
            "src/services/supabaseGateway.js 中的 request、getCatalog、getMerchants、getDishes、patchMerchant、patchDish。",
            "src/composables/useDiningStore.js 中的 loadCatalog、refreshFromCloud、queueSync。",
            "src/App.vue 中的 ensureViewData 与 watch(currentView, ...)。",
        ],
        "excerpts": [
            {"title": "统一 API 网关封装", "path": "src/services/supabaseGateway.js", "start": 1, "end": 85},
            {"title": "store 中的异步拉取与同步刷新", "path": "src/composables/useDiningStore.js", "start": 208, "end": 260},
            {"title": "顶层页面中的按视图加载入口", "path": "src/App.vue", "start": 69, "end": 92},
        ],
        "explanations": [
            {"line": "supabaseGateway 第 3 到 21 行：", "explain": "request 是统一的异步请求函数。它内部调用 fetch，并统一处理请求头、错误状态和 JSON 解析，这样业务层就不用重复写这些细节。"},
            {"line": "第 32 到 85 行：", "explain": "createSupabaseGateway 返回一个接口对象，把获取目录、获取商家、获取菜品、修改和删除等请求都封装在一起，方便前端按统一方式调用。"},
            {"line": "store 第 208 到 236 行：", "explain": "loadCatalog 是最关键的异步加载函数。它会根据 scope 决定是否真正发请求，如果该范围数据已经加载过，就直接复用已有结果。"},
            {"line": "第 220 到 225 行：", "explain": "真正发请求的地方是 const payload = await gateway.getCatalog(scope)。拿到接口结果后，会再通过 applyCatalogPayload 写回响应式状态。"},
            {"line": "第 238 到 247 行：", "explain": "refreshFromCloud 是强制刷新版本。用户点“刷新云端”时，会重新执行异步拉取，确保本地页面拿到最新数据。"},
            {"line": "第 249 到 260 行：", "explain": "queueSync 负责在本地修改后异步推送到云端，再批量刷新多个页面范围的数据。这让‘写数据’和‘重新拉取最新结果’形成闭环。"},
            {"line": "App.vue 第 69 到 71 行：", "explain": "ensureViewData 是按页面范围加载数据的入口。它不会一次性把所有页面的接口都打出来，而是根据当前视图去决定拉取哪个范围的数据。"},
            {"line": "第 84 到 92 行：", "explain": "onMounted 先加载当前页面需要的数据，watch(currentView, ...) 再在页面切换时补充拉取新视图所需数据，这就是前端懒加载的核心。"},
        ],
        "defense": "这一项可以直接对应老师的异步数据获取与展示要求：我通过 fetch 去请求接口，再把结果写回 Vue 的响应式状态，让页面自动渲染。同时我还做了按页面范围懒加载，避免首页初始化过重。",
    },
]


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    add_paragraph(
        doc,
        "现代化校园健康饮食 SPA 项目",
        size=24,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
    )
    add_paragraph(
        doc,
        "6 个功能介绍、主要函数、核心代码摘录与逐行讲解",
        size=12,
        color="555555",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
    )
    add_paragraph(
        doc,
        "说明：这 6 个功能全部明确挂靠到老师给出的 5 项考核要求，其中“动态数据渲染与视图切换”被拆成了 2 个可独立答辩的功能点。",
        size=11,
        color="555555",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
    )
    add_paragraph(
        doc,
        "适用场景：3 人小组答辩，每人可从中选择 2 个功能进行重点讲解。",
        size=11,
        color="555555",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=14,
    )

    add_heading(doc, "建议分工", level=1)
    add_number(doc, "成员 A：功能 1“顶层视图切换” + 功能 2“看板与轮盘候选池的动态筛选渲染”。")
    add_number(doc, "成员 B：功能 3“交互式数据操作” + 功能 4“表单双向绑定与校验反馈”。")
    add_number(doc, "成员 C：功能 5“浏览器原生能力集成” + 功能 6“异步数据获取与展示”。")
    add_paragraph(
        doc,
        "下面每个功能都包含：对应考核要求、功能定位、主要文件与函数、整行代码摘录、逐行讲解、答辩话术。代码摘录直接来自当前项目源码，因此你在讲解时可以和本地文件逐行对照。",
        after=10,
    )

    for index, feature in enumerate(FEATURES):
        if index > 0:
            doc.add_page_break()
        add_feature(doc, feature)

    doc.add_page_break()
    add_heading(doc, "最终说明", level=1)
    add_paragraph(
        doc,
        "如果你们后面要继续精修答辩材料，还可以在这份文档基础上继续拆成 3 份个人讲稿：每个人保留自己负责的 2 个功能，再把对应页面实际运行截图贴进去，就会更像正式答辩稿。",
    )
    add_paragraph(doc, "输出文件名：six-feature-guide.docx", size=10, color="555555")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
